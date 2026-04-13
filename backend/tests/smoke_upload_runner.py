from __future__ import annotations

import asyncio
from io import BytesIO

from docx import Document
from fastapi import UploadFile

from app.application.services.cv_parser_service import CvParserService
from app.application.services.ner_service import NerService
from app.application.use_cases.upload_cv_use_case import UploadCvUseCase


def _build_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph('Candidat Demo')
    document.add_paragraph('Compétences: Python, SQL, Machine Learning')
    document.add_paragraph('Soft skills: communication, leadership')
    document.add_paragraph('Expérience: 4 ans')
    document.add_paragraph('Formation: Master en Data Science')
    document.save(buffer)
    return buffer.getvalue()


async def main() -> None:
    file = UploadFile(filename='demo_cv.docx', file=BytesIO(_build_docx_bytes()))
    result = await UploadCvUseCase(CvParserService(), NerService()).execute([file])

    print(
        {
            'accepted': result.accepted,
            'rejected': result.rejected,
            'count': result.count,
            'parsed_candidates': len(result.parsed_candidates),
            'first_candidate_name': result.parsed_candidates[0].candidate.name if result.parsed_candidates else None,
        }
    )


if __name__ == '__main__':
    asyncio.run(main())
